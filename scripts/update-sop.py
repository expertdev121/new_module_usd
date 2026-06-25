"""Build v1.1 of the Connection Point Webhook SOP.

Writes to: C:/Users/jhavi/Documents/SOP_Connection_Point_Webhook_v1.1.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

OUT = Path(r"C:/Users/jhavi/Documents/SOP_Connection_Point_Webhook_v1.1.docx")

doc = Document()

# ---------- styles ----------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def p(text=""):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

def table_from_rows(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs:
            for run in r.runs:
                run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()

# ---------- title block ----------
title = doc.add_heading("STANDARD OPERATING PROCEDURE", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run("Connection Point Webhook: Organization Payment Registration")
sub_run.bold = True
sub_run.font.size = Pt(14)

table_from_rows(
    ["Field", "Value"],
    [
        ["Document Owner", "GiveSuite / DonorHQ Team"],
        ["Version", "1.1"],
        ["Date", "June 2026"],
        ["Classification", "Confidential – Internal Use Only"],
        ["Systems", "GoHighLevel (GHL), Connection Point, DonorHQ"],
    ],
)

# ---------- 1 Purpose ----------
h1("1.  Purpose")
p(
    "This Standard Operating Procedure (SOP) defines the end-to-end process for handling "
    "incoming donation and payment webhook events originating from Connection Point and routing "
    "them through the GoHighLevel (GHL) automation platform. The workflow automatically creates "
    "or updates donor contacts, enriches them with campaign and payment data, and forwards the "
    "relevant donation information to DonorHQ for record-keeping and downstream processing."
)
p(
    "It also documents the behavior of the DonorHQ inbound webhook endpoint "
    "(/api/webhook/ghl/donation) — what happens to the payload once it reaches DonorHQ, how "
    "contacts and campaigns are reconciled, how duplicates are prevented, and what is returned to "
    "GHL on each outcome."
)
p(
    "This document is intended to serve as the authoritative reference for configuring, "
    "maintaining, troubleshooting, and testing this integration."
)

# ---------- 2 Scope ----------
h1("2.  Scope")
p("This SOP applies to:")
bullet("All payment.created webhook events triggered by Connection Point for organizational entities.")
bullet('The GoHighLevel workflow named "Connection Point Webhook" (published state).')
bullet("Integration personnel responsible for configuring or maintaining GHL automation workflows.")
bullet("Development and QA teams testing donation flows between Connection Point, GHL, and DonorHQ.")
bullet("The DonorHQ inbound webhook endpoint that accepts the forwarded donation data.")
p(
    "Out of scope: individual donor-level webhooks not classified as entity_type: organization, and "
    "any billing or refund events not covered by the payment.created event type."
)

# ---------- 3 Tools & Systems ----------
h1("3.  Tools & Systems")
table_from_rows(
    ["System", "Purpose"],
    [
        ["Connection Point", "Donation platform that fires webhook events on payment creation"],
        ["GoHighLevel (GHL)", "Automation platform that receives, processes, and routes webhook data"],
        ["DonorHQ", "Donation management system receiving enriched donor & payment data"],
        ["Stripe", "Underlying payment processor referenced in the transaction_id"],
        ["GHL Inbound Webhook", "Entry trigger for the GHL workflow (Connection Point data)"],
        ["GHL Custom Webhook Action", "Outbound POST to DonorHQ endpoint at donorhq.givesuite.com"],
    ],
)

# ---------- 4 Trigger / Entry Point ----------
h1("4.  Trigger / Entry Point")
p(
    "The GHL workflow is triggered by an Inbound Webhook event sent from Connection Point whenever "
    "a payment is successfully created. The trigger is configured with no filters, meaning all "
    "incoming payloads matching the webhook URL will initiate the workflow."
)

h2("4.1  Webhook URLs")
p("The following GHL inbound webhook URLs are registered to receive Connection Point payment events:")
table_from_rows(
    ["Label", "URL"],
    [
        [
            "Primary",
            "https://services.leadconnectorhq.com/hooks/sNXq6gyPrArxiSrFEaaf/webhook-trigger/EJpVL31yqAtRWpYoX6Ki",
        ],
        [
            "Secondary",
            "https://services.leadconnectorhq.com/hooks/KVgMIrEYRkKRcfeicJBm/webhook-trigger/4I22l6Wbnt9LpLs4khv2",
        ],
    ],
)
p(
    "Note: The first path segment after /hooks/ is the GHL location_id (e.g. sNXq6gyPrArxiSrFEaaf "
    "for the Primary URL). This value MUST be forwarded to DonorHQ as the location_id custom-data "
    "field (see Section 5, Step 6) so that DonorHQ can scope contact and campaign lookups correctly. "
    "As a fallback, the DonorHQ endpoint also accepts location_id as a URL query parameter — see "
    "Section 7.1 — so a misconfigured workflow can still be recovered with a one-character URL edit."
)

h2("4.2  Trigger Payload Fields")
p("Key fields extracted from the inbound webhook payload:")
table_from_rows(
    ["Field", "Description"],
    [
        ["event", "Event type — always payment.created for this flow"],
        ["entity_type", "Entity classification — organization"],
        ["entity_id", "Unique ID of the organization"],
        ["data.payer_email", "Email address of the donor/payer"],
        ["data.payer_first_name", "First name of the donor/payer"],
        ["data.payer_last_name", "Last name of the donor/payer"],
        ["data.amount", "Gross donation amount (e.g., 101)"],
        ["data.net_amount", "Net amount after fees (e.g., 94.74)"],
        ["data.transaction_id", "Stripe subscription/transaction ID — used as dedup key in DonorHQ"],
        ["data.campaign_id", "Connection Point campaign ID"],
        ["data.organization_id", "Organization that received the donation"],
        ["data.payment_type", "Payment type — e.g., subscription, one_time"],
        ["data.status", "Payment status — e.g., preapproved, completed"],
        ["data.currency", "Currency code — e.g., usd"],
        ["data.account", "Payment processor account — e.g., stripe"],
        ["data.created", "Unix timestamp of payment creation — forwarded as received_date"],
    ],
)

# ---------- 5 Step-by-Step Workflow ----------
h1("5.  Step-by-Step Workflow")
p('The GHL workflow "Connection Point Webhook" consists of the following steps executed in sequence:')

h3("Step 1 — Inbound Webhook Trigger (Connection Point data)")
bullet("Trigger type: Inbound Webhook")
bullet("Trigger name: Connection Point data")
bullet("Filters: None applied")
bullet("Fires when: A POST request is received at the registered webhook URL")
bullet("Event condition: event = payment.created")
bullet("Entity condition: entity_type = organization")

h3("Step 2 — Create / Update Contact")
bullet("GHL Action: Create Contact")
bullet("Behavior: Creates a new contact if none exists; updates if email already on file")
bullet("Field mappings:")
bullet("    Email → {{inboundWebhookRequest.data.payer_email}}")
bullet("    First Name → {{inboundWebhookRequest.data.payer_first_name}}")
bullet("    Last Name → {{inboundWebhookRequest.data.payer_last_name}}")
bullet("Note: Only works with contactless execution methods (Inbound Webhook)")

h3("Step 3 — Add Tag")
bullet("GHL Action: Add Tag")
bullet("Purpose: Tags the contact for segmentation, filtering, and reporting")
bullet("Configure tag value to reflect payment source, campaign, or status as needed")

h3("Step 4 — #1 Fetch Campaign Details")
bullet("GHL Action: Custom Value / API Lookup")
bullet("Purpose: Retrieves campaign-level metadata associated with the campaign_id")
bullet("Input: {{inboundWebhookRequest.data.campaign_id}}")
bullet("Output: Campaign name (used in Step 6 as campaign_name), description, organization details")

h3("Step 5 — #2 Fetch Payment Details")
bullet("GHL Action: Custom Value / API Lookup")
bullet("Purpose: Retrieves complete payment record for the transaction")
bullet("Input: {{inboundWebhookRequest.data.transaction_id}}")
bullet("Output: Full payment object including fees, net amount, status, and metadata")

h3("Step 6 — Send Data To DonorHq")
bullet("GHL Action: Webhook (outbound POST)")
bullet("Method: POST")
bullet("URL: https://donorhq.givesuite.com/api/webhook/ghl/donation")
bullet("Content-Type: application/json")
bullet("Custom data sent (all wired in the workflow Custom Data section):")
table_from_rows(
    ["DonorHQ Field", "GHL Mapped Value", "Required?"],
    [
        ["location_id", "sNXq6gyPrArxiSrFEaaf (hardcoded — matches inbound URL location segment)", "Yes"],
        ["contact_id", "{{contact.id}}", "Yes (or email)"],
        ["email", "{{inboundWebhookRequest.data.payer_email}}", "Yes (or contact_id)"],
        ["first_name", "{{inboundWebhookRequest.data.payer_first_name}}", "Recommended"],
        ["last_name", "{{inboundWebhookRequest.data.payer_last_name}}", "Recommended"],
        ["campaign_name", "{{custom_webhook.1.response.title}} (from Step 4)", "Yes"],
        ["campaign_id", "{{inboundWebhookRequest.data.campaign_id}}", "Optional"],
        ["donation_amount", "{{inboundWebhookRequest.data.amount}}", "Yes"],
        ["currency", "{{inboundWebhookRequest.data.currency}}", "Optional (defaults to USD)"],
        ["transaction_id", "{{inboundWebhookRequest.data.transaction_id}}", "Recommended — drives dedup"],
        ["payment_type", "{{inboundWebhookRequest.data.payment_type}}", "Optional"],
        ["received_date", "{{inboundWebhookRequest.data.created}} or workflow event date", "Optional (defaults to today)"],
    ],
)

h3("Step 7 — END")
bullet("Workflow execution completes.")
bullet("Contact is created/updated in GHL.")
bullet("Donation data has been forwarded to DonorHQ.")

# ---------- 6 Field Mapping Reference ----------
h1("6.  Field Mapping Reference")

h2("6.1  GHL Contact Fields (Step 2)")
table_from_rows(
    ["GHL Field", "Mapped Value"],
    [
        ["Email", "{{inboundWebhookRequest.data.payer_email}}"],
        ["First Name", "{{inboundWebhookRequest.data.payer_first_name}}"],
        ["Last Name", "{{inboundWebhookRequest.data.payer_last_name}}"],
    ],
)

h2("6.2  DonorHQ Custom Data Fields (Step 6)")
p(
    "All custom-data fields configured in the Send Data To DonorHq webhook action. "
    "Required fields will cause DonorHQ to return 400 with a specific error code if missing."
)
table_from_rows(
    ["DonorHQ Field", "GHL Mapped Value", "Notes"],
    [
        ["location_id", "sNXq6gyPrArxiSrFEaaf", "Hardcoded; matches inbound URL location segment. Scopes all DonorHQ lookups."],
        ["contact_id", "{{contact.id}}", "GHL's contact ID; used as ghl_contact_id in DonorHQ."],
        ["email", "{{inboundWebhookRequest.data.payer_email}}", "Fallback contact-lookup key, case-insensitive."],
        ["first_name", "{{inboundWebhookRequest.data.payer_first_name}}", "Used only when DonorHQ creates a new contact."],
        ["last_name", "{{inboundWebhookRequest.data.payer_last_name}}", "Used only when DonorHQ creates a new contact."],
        ["campaign_name", "{{custom_webhook.1.response.title}}", "Find-or-create within the location."],
        ["campaign_id", "{{inboundWebhookRequest.data.campaign_id}}", "Recorded in donation notes for traceability."],
        ["donation_amount", "{{inboundWebhookRequest.data.amount}}", "Must be > 0."],
        ["currency", "{{inboundWebhookRequest.data.currency}}", "Defaults to USD if absent. Supported: USD/ILS/EUR/JPY/GBP/AUD/CAD/ZAR."],
        ["transaction_id", "{{inboundWebhookRequest.data.transaction_id}}", "Drives dedup via (location_id, transaction_id)."],
        ["payment_type", "{{inboundWebhookRequest.data.payment_type}}", "Stored as payment_method (defaults to 'unknown')."],
        ["received_date", "{{inboundWebhookRequest.data.created}}", "Used as both payment_date and received_date. Defaults to today."],
    ],
)

# ---------- 7 DonorHQ Endpoint Behavior (NEW) ----------
h1("7.  DonorHQ Webhook Endpoint Behavior")
p(
    "This section documents exactly what happens server-side once the GHL workflow POSTs to "
    "DonorHQ. It is the canonical reference for support and engineering when debugging payloads "
    "or unexpected outcomes."
)

h2("7.1  Endpoint")
table_from_rows(
    ["Attribute", "Value"],
    [
        ["URL", "POST https://donorhq.givesuite.com/api/webhook/ghl/donation\n(optional URL query fallback: ?location_id=<id> — used only if the body does not include location_id)"],
        ["Source file", "app/api/webhook/ghl/donation/route.ts"],
        ["Authentication", "Public — excluded from the NextAuth matcher (middleware.ts). No bearer token, no signature check."],
        ["Accepted content types", "application/json, application/x-www-form-urlencoded, multipart/form-data"],
        ["Idempotency", "Driven by (location_id, transaction_id). Repeat posts return 200 DONATION_ALREADY_EXISTS with the existing row."],
        ["Correlation ID", "Every response and log line includes a UUID reqId for tracing a single request end-to-end."],
    ],
)

h2("7.2  Processing Pipeline")
p("In order, the endpoint executes the following steps. Failure at any step is logged with reqId and returned as a structured JSON error.")

h3("Step 1 — Body parse")
bullet("Reads the request body based on Content-Type (JSON, form-encoded, or raw text).")
bullet("On parse failure: returns 400 BODY_PARSE_ERROR with the raw error message.")

h3("Step 2 — Flatten GHL envelope")
p(
    "GHL workflow webhooks wrap workflow-configured custom fields inside a customData object "
    "alongside the standard contact data. The endpoint flattens this so the validator can see the "
    "fields at the top level. Three shapes are handled:"
)
bullet("customData as an object → keys merged to top level")
bullet("customData as a JSON string → parsed, then merged")
bullet("customData[key]= form-encoded entries → bracket-stripped and lifted")
p(
    "Fallback: if location_id is still missing after flattening, the endpoint inspects the standard "
    "location field (often a {id, name} object or stringified version) and extracts .id."
)

h3("Step 3 — Schema validation")
p("Required fields:")
bullet("location_id (or locationId) — also accepted as ?location_id=… in the URL; body wins on conflict")
bullet("campaign_name (or campaignName)")
bullet("donation_amount (or amount) — must parse to a positive number")
bullet("At least one of contact_id or email")
p("Errors returned (HTTP 400):")
bullet("VALIDATION_ERROR — Zod schema failure; includes per-field details")
bullet("MISSING_LOCATION_ID — location_id absent or empty")
bullet("MISSING_CONTACT_IDENTIFIER — neither contact_id nor email provided")
bullet("MISSING_CAMPAIGN_NAME — campaign_name absent or empty")
bullet("INVALID_AMOUNT — donation_amount missing, non-numeric, or non-positive")
bullet("INVALID_RECEIVED_DATE — received_date present but not parseable (YYYY-MM-DD or ISO timestamp)")

h3("Step 4 — Contact find-or-create (scoped by location_id)")
p("Lookup order, all scoped to the provided location_id:")
bullet("By (location_id, ghl_contact_id = contact_id)")
bullet("Fallback by (location_id, email) — case-insensitive")
p(
    "If no match is found, a new contact row is created with the supplied first_name / last_name / "
    "email / contact_id / location_id. Missing first or last names fall back to \"N/A\" so the "
    "NOT NULL columns are satisfied. display_name is set to the joined first + last (with N/A "
    "stripped), or null if both names fell back."
)
p("On DB error: returns 500 CONTACT_UPSERT_FAILED.")
p("Logs: contact-created (info) when a new row is written.")

h3("Step 5 — Campaign find-or-create (scoped by location_id)")
p("Lookup by (campaign.name = campaign_name, campaign.location_id = location_id).")
p("If no match, a new campaign row is created with status = active.")
p("On DB error: returns 500 CAMPAIGN_UPSERT_FAILED.")
p("Logs: campaign-created (info) when a new row is written.")

h3("Step 6 — Manual donation insert (with dedup)")
p("Inserts a row into manual_donation with the following columns populated:")
bullet("contact_id, campaign_id, location_id")
bullet("amount (string, 2 decimals), currency, amount_usd, exchange_rate")
bullet("payment_date and received_date (both set to received_date if provided, else today)")
bullet("payment_method (= payment_type), payment_status = completed")
bullet("ghl_source = 'ghl_workflow', ghl_resource_id = transaction_id, ghl_payment_method = payment_type")
bullet("reference_number = transaction_id")
bullet("notes (auto-built): includes payer name, transaction_id, ghl_campaign_id, sub_campaign_id when present")
p(
    "Dedup: when transaction_id is present, the insert uses ON CONFLICT DO NOTHING against the "
    "partial unique index manual_donation_ghl_location_unique on (location_id, ghl_resource_id). "
    "If the insert returns no row, the existing row is fetched and returned with code "
    "DONATION_ALREADY_EXISTS (HTTP 200)."
)
p("On DB error: returns 500 DONATION_INSERT_FAILED.")

h2("7.3  Response Codes")
table_from_rows(
    ["HTTP", "Code", "Meaning"],
    [
        ["201", "DONATION_CREATED", "New donation row inserted successfully."],
        ["200", "DONATION_ALREADY_EXISTS", "A donation with the same (location_id, transaction_id) already existed; existing row returned."],
        ["400", "BODY_PARSE_ERROR", "Request body could not be parsed."],
        ["400", "VALIDATION_ERROR", "Zod schema validation failed; per-field issues included."],
        ["400", "MISSING_LOCATION_ID", "location_id missing or empty after flattening."],
        ["400", "MISSING_CONTACT_IDENTIFIER", "Neither contact_id nor email provided."],
        ["400", "MISSING_CAMPAIGN_NAME", "campaign_name missing or empty."],
        ["400", "INVALID_AMOUNT", "donation_amount missing, non-numeric, or not positive."],
        ["400", "INVALID_RECEIVED_DATE", "received_date is not parseable as a date."],
        ["500", "CONTACT_UPSERT_FAILED", "Database error during contact lookup or insert."],
        ["500", "CAMPAIGN_UPSERT_FAILED", "Database error during campaign lookup or insert."],
        ["500", "DONATION_INSERT_FAILED", "Database error during donation insert (or insert returned no row)."],
        ["500", "SERVER_ERROR", "Unhandled exception; stack returned in non-production."],
    ],
)

h2("7.4  Success Response Shape")
code(
    "{\n"
    '  "success": true,\n'
    '  "code": "DONATION_CREATED",\n'
    '  "reqId": "9597d0df-...",\n'
    '  "data": {\n'
    '    "donation": { ... full manual_donation row ... },\n'
    '    "contact": { "id": 123, "firstName": "Orna", "lastName": "beer", "created": false },\n'
    '    "campaign": { "id": 45, "name": "For Parnassah Tova!", "created": true },\n'
    '    "locationId": "sNXq6gyPrArxiSrFEaaf",\n'
    '    "transactionId": "sub_1TlWCjFVb9BFQ9XAcgNvlgwr"\n'
    "  }\n"
    "}"
)

h2("7.5  Error Response Shape")
code(
    "{\n"
    '  "success": false,\n'
    '  "code": "MISSING_LOCATION_ID",\n'
    '  "message": "location_id is required",\n'
    '  "reqId": "9597d0df-...",\n'
    '  "receivedKeys": ["email", "contact_id", ...],   // post-flatten keys\n'
    '  "rawKeys":      ["email", "customData", ...]    // envelope keys\n'
    "}"
)

h2("7.6  Logging")
p('All log lines are prefixed with "[ghl-donation-webhook] <reqId>" so a single request can be traced end-to-end.')
table_from_rows(
    ["Level", "Event", "Triggered when"],
    [
        ["log", "incoming", "Every request, before validation. Includes Content-Type and raw body."],
        ["log", "flattened customData → keys=…", "When customData was present and merged."],
        ["log", "contact-created", "A new contact row was inserted."],
        ["log", "campaign-created", "A new campaign row was inserted."],
        ["log", "success / duplicate", "Donation insert completed; includes donation/contact/campaign IDs, amount, currency, elapsed ms."],
        ["warn", "fail status=… code=…", "Any 4xx/5xx error returned to the client."],
        ["error", "<step>-failed", "Database errors or unhandled exceptions; includes the full error and inputs."],
    ],
)

h2("7.7  Idempotency Guarantee")
p(
    "Because GHL can retry webhook deliveries, the endpoint is safe to retry. When transaction_id "
    "is included in the payload, repeated posts for the same (location_id, transaction_id) "
    "return 200 DONATION_ALREADY_EXISTS with the original donation row, instead of creating "
    "duplicates. If transaction_id is NOT included, dedup is skipped and each retry creates a new "
    "donation row — so always wire transaction_id in the GHL workflow."
)

# ---------- 8 Sample Webhook Payload ----------
h1("8.  Sample Webhook Payload")

h2("8.1  Connection Point → GHL Inbound")
p("Representative Connection Point payment.created payload used for testing and validation:")
code(
    "{\n"
    '  "event": "payment.created",\n'
    '  "entity_type": "organization",\n'
    '  "entity_id": "4JMJ1",\n'
    '  "data": {\n'
    '    "created": 1782228334,\n'
    '    "campaign_id": "e2iyZ4",\n'
    '    "organization_id": "4JMJ1",\n'
    '    "amount": 101,\n'
    '    "net_amount": 94.74,\n'
    '    "currency": "usd",\n'
    '    "status": "preapproved",\n'
    '    "payer_name": "Orna beer",\n'
    '    "payer_first_name": "Orna",\n'
    '    "payer_last_name": "beer",\n'
    '    "payer_email": "orpro1976@gmail.com",\n'
    '    "transaction_id": "sub_1TlWCjFVb9BFQ9XAcgNvlgwr",\n'
    '    "account": "stripe",\n'
    '    "payment_type": "subscription",\n'
    '    "id": "7769331",\n'
    '    "object": "payment"\n'
    "  },\n"
    '  "id": "3sC",\n'
    '  "object": "webhook_event"\n'
    "}"
)

h2("8.2  GHL → DonorHQ Outbound (Step 6)")
p('Representative payload posted by the GHL "Send Data To DonorHq" webhook action:')
code(
    "{\n"
    '  "location_id": "sNXq6gyPrArxiSrFEaaf",\n'
    '  "contact_id": "abc123-ghl-contact-id",\n'
    '  "email": "orpro1976@gmail.com",\n'
    '  "first_name": "Orna",\n'
    '  "last_name": "beer",\n'
    '  "campaign_name": "For Parnassah Tova!",\n'
    '  "campaign_id": "e2iyZ4",\n'
    '  "donation_amount": "101",\n'
    '  "currency": "usd",\n'
    '  "transaction_id": "sub_1TlWCjFVb9BFQ9XAcgNvlgwr",\n'
    '  "payment_type": "subscription",\n'
    '  "received_date": "2026-06-25"\n'
    "}"
)

h2("8.3  Manual Test")
p("To test the inbound webhook manually:")
code(
    'curl -X POST "https://services.leadconnectorhq.com/hooks/sNXq6gyPrArxiSrFEaaf/webhook-trigger/EJpVL31yqAtRWpYoX6Ki" \\\n'
    '  -H "Content-Type: application/json" \\\n'
    "  -d @payload.json"
)
p("To test the DonorHQ endpoint directly (skipping GHL):")
code(
    'curl -X POST "https://donorhq.givesuite.com/api/webhook/ghl/donation" \\\n'
    '  -H "Content-Type: application/json" \\\n'
    "  -d @donorhq_payload.json"
)

# ---------- 9 Workflow Diagram ----------
h1("9.  Workflow Diagram")
p(
    'The following screenshots are taken directly from the GoHighLevel Workflow Builder and '
    'illustrate the live configuration of the "Connection Point Webhook" automation.'
)

h2("9.1  Full Workflow Builder — Create Contact Action")
p("[ Screenshot 1: GHL Workflow Builder — Create Contact Action ]")
p(
    "This screenshot shows the full workflow canvas with steps: Inbound Webhook → Create Contact → "
    "Add Tag → Fetch Campaign Details → Fetch Payment Details → Send Data To DonorHq → END. The "
    "right panel shows the Create Contact action with Email, First Name, and Last Name field "
    "mappings."
)

h2("9.2  Send Data To DonorHq — Webhook Action Configuration")
p("[ Screenshot 2: GHL Webhook Action — Send Data To DonorHq ]")
p(
    "This screenshot shows the outbound webhook configuration: Method: POST | URL: "
    "https://donorhq.givesuite.com/api/webhook/ghl/donation | Custom data fields: location_id, "
    "contact_id, email, first_name, last_name, campaign_name, campaign_id, donation_amount, "
    "currency, transaction_id, payment_type, received_date."
)

# ---------- 10 Notes & Troubleshooting ----------
h1("10.  Notes & Troubleshooting")

h2("10.1  General Notes")
bullet("The GHL workflow must be in Published state for live webhook events to trigger execution. Draft mode will not process inbound events.")
bullet('The "Create Contact" action only works with contactless execution methods. The inbound webhook satisfies this requirement.')
bullet("If a contact with the payer_email already exists in GHL, the existing record will be updated rather than duplicated.")
bullet("The Add Tag step should be configured with a meaningful tag value (e.g., connection-point-donor, subscription-payment) to enable effective segmentation.")
bullet("Both webhook URLs (Primary and Secondary) accept the same payload format and trigger the same workflow logic.")
bullet("location_id MUST be set to the location segment from the inbound URL (sNXq6gyPrArxiSrFEaaf for Primary). Do not derive it from the payload.")
bullet("transaction_id MUST be wired in Step 6 to prevent duplicate donations on webhook retries.")

h2("10.2  Common Issues & Resolutions")
table_from_rows(
    ["Issue", "Resolution"],
    [
        ["Webhook returns 404", "Verify the webhook URL is correct and the GHL workflow is published, not in Draft mode."],
        ["Contact not created in GHL", "Check that payer_email is present in the payload. Confirm the Create Contact step is active and properly mapped."],
        ["DonorHQ not receiving data", "Verify the outbound webhook URL in Step 6 is correct. Check DonorHQ server logs for incoming requests. Confirm GHL workflow reaches Step 6 via Execution Logs."],
        ["DonorHQ returns 400 MISSING_LOCATION_ID", "Add a location_id custom-data field in Step 6 with the hardcoded value sNXq6gyPrArxiSrFEaaf, OR append ?location_id=sNXq6gyPrArxiSrFEaaf to the outbound URL in Step 6."],
        ["DonorHQ returns 400 MISSING_CAMPAIGN_NAME", "Add a campaign_name custom-data field in Step 6 mapped to the Step 4 campaign lookup output (e.g. {{custom_webhook.1.response.title}})."],
        ["DonorHQ returns 400 INVALID_AMOUNT", "Check that donation_amount is mapped to {{inboundWebhookRequest.data.amount}} and that the Connection Point payload includes a positive number."],
        ["DonorHQ returns 400 INVALID_RECEIVED_DATE", "Verify the received_date mapping resolves to a valid date string. Unix timestamps are accepted via the ISO date parser."],
        ["DonorHQ returns 200 DONATION_ALREADY_EXISTS", "Expected — a previous post with the same (location_id, transaction_id) already created the donation. No action required."],
        ["Duplicate contacts created in GHL", "Ensure email field mapping uses data.payer_email (not payer_name or another field). GHL deduplicates by email."],
        ["Duplicate donations in DonorHQ", "Verify transaction_id is wired in Step 6. Without it, dedup is disabled."],
        ["Workflow not triggering", 'Confirm the workflow is Published. Test using the GHL "Test workflow" button or send a manual curl request to the webhook URL.'],
        ["Fields are empty in DonorHQ", "Validate that the inbound payload contains the expected fields. Use GHL Execution Logs to inspect the raw inbound webhook data."],
    ],
)

h2("10.3  Where to Check Execution Logs")
bullet("GHL: Automation > Workflows > Connection Point Webhook > Execution Logs tab. Each execution shows the triggered contact, step-by-step status, and any errors encountered.")
bullet("GHL: Enrollment History tab to see all contacts that have entered this workflow.")
bullet("DonorHQ: Server logs filtered by prefix [ghl-donation-webhook]. Every line is tagged with the per-request reqId from the response body — grep for it to see the full lifecycle of a single request.")
bullet("DonorHQ: The 'incoming' log line includes the raw posted body (verbatim JSON), useful when the GHL Execution Log only shows the field mappings rather than the final HTTP body.")

# ---------- 11 Revision History ----------
h1("11.  Revision History")
table_from_rows(
    ["Version", "Date", "Author", "Notes"],
    [
        ["1.0", "June 2026", "GiveSuite / DonorHQ Team", "Initial release."],
        ["1.1", "June 2026", "GiveSuite / DonorHQ Team", "Expanded Step 6 custom-data field list (location_id, campaign_name, currency, payment_type, received_date, etc.). Added Section 7 documenting the DonorHQ endpoint behavior, response codes, logging, and idempotency. Documented URL query-param fallback for location_id. Expanded troubleshooting table with DonorHQ-side error codes."],
    ],
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Wrote {OUT}")
